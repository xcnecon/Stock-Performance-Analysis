"""Render results/report.docx from the calendar-window analysis outputs.

Reads results/summary_{10y,30y,fulllife}.csv, market_benchmark.csv,
market_summary.csv, and the chart PNGs produced by run_analysis.py.
"""
from __future__ import annotations

import os
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = r"C:\Users\CHENY\Documents\GitHub\Stock-Performance-Analysis"
RESULTS = os.path.join(ROOT, "results")


# ---------- cell / table helpers ----------
def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              header_fill: str = "E2E8F0",
              col_widths: Sequence[float] | None = None) -> None:
    rows = list(rows)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True

    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, header_fill)

    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(v))
            run.font.size = Pt(9)

    if col_widths:
        for row in tbl.rows:
            for j, width_in in enumerate(col_widths):
                row.cells[j].width = Inches(width_in)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def add_picture(doc: Document, filename: str, width_in: float = 6.8) -> None:
    path = os.path.join(RESULTS, filename)
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------- formatters ----------
def pct(x: float, d: int = 1) -> str:
    return f"{x * 100:.{d}f} %"


def signed_pct(x: float, d: int = 1) -> str:
    if pd.isna(x):
        return "—"
    sign = "+" if x >= 0 else "-"
    return f"{sign}{abs(x) * 100:.{d}f} %"


def fmt_mult(x: float) -> str:
    if pd.isna(x):
        return "—"
    if x < 0:
        return f"({x * 100:.1f} %)"
    if 1 + x < 10:
        return f"{(1+x):,.2f}×"
    return f"{(1+x):,.0f}×"


# ---------- build doc ----------
def build(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Long-run return distributions of US-listed stocks")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("CRSP monthly, 31,565 common stocks, Dec 1925 – Dec 2025")
    rr.italic = True
    rr.font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub2.add_run("Calendar-window methodology: every 10-year and 30-year window analyzed on its own cohort")
    rr.italic = True
    rr.font.size = Pt(10)

    doc.add_paragraph()

    # Load data
    s10 = pd.read_csv(os.path.join(RESULTS, "summary_10y.csv"))
    s30 = pd.read_csv(os.path.join(RESULTS, "summary_30y.csv"))
    s_life = pd.read_csv(os.path.join(RESULTS, "summary_fulllife.csv"))
    bench = pd.read_csv(os.path.join(RESULTS, "market_benchmark.csv"))
    mkt = pd.read_csv(os.path.join(RESULTS, "market_summary.csv")).iloc[0].to_dict()

    bench10 = bench[bench["horizon"] == "10y"].reset_index(drop=True)
    bench30 = bench[bench["horizon"] == "30y"].reset_index(drop=True)
    bench_life = bench[bench["horizon"] == "full-life"].iloc[0]

    # ============= TL;DR =============
    add_heading(doc, "TL;DR", 1)

    worst_nom_10 = s10[s10["basis"] == "nominal"].nsmallest(1, "pct_positive").iloc[0]
    best_nom_10 = s10[s10["basis"] == "nominal"].nlargest(1, "pct_positive").iloc[0]
    worst_real_10 = s10[s10["basis"] == "real (CPI-adj)"].nsmallest(1, "pct_positive").iloc[0]
    best_real_10 = s10[s10["basis"] == "real (CPI-adj)"].nlargest(1, "pct_positive").iloc[0]

    b = doc.add_paragraph(style="List Bullet")
    b.add_run("The share of stocks with a positive return is ")
    b.add_run("deeply period-dependent").bold = True
    b.add_run(f". Across the ten non-overlapping decades since 1925, % positive (nominal) ranges from ")
    b.add_run(pct(worst_nom_10['pct_positive'])).bold = True
    b.add_run(f" ({worst_nom_10['window']}) to ")
    b.add_run(pct(best_nom_10['pct_positive'])).bold = True
    b.add_run(f" ({best_nom_10['window']}).")

    b = doc.add_paragraph(style="List Bullet")
    b.add_run("In real (CPI-adjusted) terms, the worst decade is ")
    b.add_run(worst_real_10['window']).bold = True
    b.add_run(f" — just ")
    b.add_run(pct(worst_real_10['pct_positive'])).bold = True
    b.add_run(" of stocks produced a positive real return as stagflation eroded nominal gains.")

    # concentration finding
    conc = bench10[bench10["pct_beat_vwretd"] < 0.20].sort_values("start_ymm")
    concs = ", ".join(f"{row['window']} ({pct(row['pct_beat_vwretd'])})" for _, row in conc.iterrows())
    if concs:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run("The market return and the share of individual stocks beating it ")
        b.add_run("diverge sharply in mega-cap-concentrated bull decades").bold = True
        b.add_run(f": {concs}.")

    s30_nom = s30[s30["basis"] == "nominal"]
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Over 30-year windows, ")
    b.add_run(f"{pct(s30_nom['pct_positive'].min())} – {pct(s30_nom['pct_positive'].max())}").bold = True
    b.add_run(" of stocks produced a positive nominal return. The share fell from ")
    b.add_run(pct(s30_nom.iloc[1]['pct_positive'])).bold = True
    b.add_run(" in 1955-1985 to ")
    b.add_run(pct(s30_nom.iloc[2]['pct_positive'])).bold = True
    b.add_run(" in 1985-2015.")

    life_nom = s_life[s_life["basis"] == "nominal"].iloc[0]
    life_real = s_life[s_life["basis"] == "real (CPI-adj)"].iloc[0]
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Over each stock's entire listed life (Q3), only ")
    b.add_run(pct(life_nom['pct_positive'])).bold = True
    b.add_run(" produced a positive nominal return and ")
    b.add_run(pct(life_real['pct_positive'])).bold = True
    b.add_run(" produced a positive real return.")

    # ============= 1. Data & universe =============
    add_heading(doc, "1. Data & universe", 1)
    add_para(doc,
             f"Source: CRSP Monthly Stock File (data/data.csv, 5,307,138 rows, 40,499 raw PERMNOs). "
             f"CPI-U All Urban Consumers, not seasonally adjusted (BLS monthly, 1925-2025), cached at data/cpi.csv. "
             f"Full-span cumulative inflation {(1+mkt['vwretd_total_return'])/(1+mkt['vwretd_total_return']):.1f} — correction: "
             f"18.1× (2.94 % / yr).")
    p = doc.add_paragraph()
    p.add_run("Universe: ").bold = True
    p.add_run("SecurityType == 'EQTY' AND IssuerType != 'REIT' — 31,565 common stocks (22,623 CORP + 8,867 ACOR including 1,377 ADRs). "
              "Excludes closed-end funds & ETFs (classified as 'FUND' in CRSP), derivatives ('DERV'), and REITs.")
    add_para(doc, "Span: 1925-12 – 2025-12 (1,201 months).")

    # ============= 2. Methodology =============
    add_heading(doc, "2. Methodology — calendar-aligned windows", 1)
    add_para(doc,
             "Per CLAUDE.md the analysis must be survivorship-bias-free: every stock that ever listed is retained, "
             "with delisting handled by terminating the return stream at the delisting month "
             "(CRSP embeds the delisting return in MthRet when MthRetFlg == 'DE').")

    p = doc.add_paragraph()
    p.add_run("For each calendar window (T, T+N]:").bold = True
    doc.add_paragraph(
        "Cohort = every common stock with a row at month T (listed at end of month T). Stocks that IPO after T are not in this window's cohort.",
        style="List Number")
    doc.add_paragraph(
        "Window return = ∏(1 + MthRet) − 1 compounded over months strictly after T and up to T+N, truncated at the stock's last listed month.",
        style="List Number")
    doc.add_paragraph(
        "Partial windows (stock delists during the window) are kept, not dropped.",
        style="List Number")

    p = doc.add_paragraph()
    p.add_run("Windows used:").bold = True
    add_table(doc,
              ["Horizon", "Windows", "Months each"],
              [
                  ["10-year", "1925-12 → 1935-12, 1935-12 → 1945-12, …, 2015-12 → 2025-12 (10 decades)", "120"],
                  ["30-year", "1925-12 → 1955-12, 1955-12 → 1985-12, 1985-12 → 2015-12 (3 eras)", "360"],
                  ["Full-life", "per-stock: first listed month → last listed month", "variable"],
              ],
              col_widths=[1.0, 4.8, 1.0])
    doc.add_paragraph()

    add_para(doc,
             "Real returns: monthly real return (1 + MthRet) / (1 + π) − 1, where π = CPI_t / CPI_{t-1} − 1, "
             "then compounded the same way as nominal.")

    p = doc.add_paragraph()
    p.add_run("Benchmark:").bold = True
    doc.add_paragraph(
        "CRSP value-weighted total return (vwretd) — with dividends; apples-to-apples with MthRet.",
        style="List Bullet")
    doc.add_paragraph(
        "S&P 500 price index (sprtrn) — no dividends; reported for reference.",
        style="List Bullet")
    doc.add_paragraph(
        "Both are compounded over the same calendar months as the window. '% beating market' = share of the cohort whose nominal return exceeded the benchmark.",
        style="List Bullet")

    # ============= 3. Q1 10-year =============
    add_heading(doc, "3. Q1 — 10-year calendar windows", 1)

    add_heading(doc, "Share of stocks positive, and % beating market, decade by decade", 2)
    rows10 = []
    for _, row in bench10.iterrows():
        nom = s10[(s10["window"] == row["window"]) & (s10["basis"] == "nominal")].iloc[0]
        real = s10[(s10["window"] == row["window"]) & (s10["basis"] == "real (CPI-adj)")].iloc[0]
        rows10.append([
            row["window"],
            f"{int(row['n_stocks']):,}",
            pct(nom["pct_positive"]),
            pct(real["pct_positive"]),
            signed_pct(row["vwretd_window_return"], 0),
            pct(row["pct_beat_vwretd"]),
        ])
    add_table(doc,
              ["Decade", "N", "% positive nominal", "% positive real", "Market return (vwretd)", "% beating market"],
              rows10,
              col_widths=[1.1, 0.8, 1.3, 1.2, 1.6, 1.3])
    doc.add_paragraph()

    add_picture(doc, "variation_10y.png")
    add_picture(doc, "pct_positive_10y.png")

    add_heading(doc, "Medians and percentiles", 2)
    rows_pct = []
    for _, row in bench10.iterrows():
        for basis_label, basis_key in [("Nominal", "nominal"), ("Real", "real (CPI-adj)")]:
            s = s10[(s10["window"] == row["window"]) & (s10["basis"] == basis_key)].iloc[0]
            rows_pct.append([
                row["window"], basis_label,
                signed_pct(s["p5"], 1), signed_pct(s["p25"], 1),
                signed_pct(s["p50"], 1), signed_pct(s["p75"], 1), signed_pct(s["p95"], 1),
            ])
    add_table(doc,
              ["Decade", "Basis", "p5", "p25", "Median", "p75", "p95"],
              rows_pct,
              col_widths=[1.0, 0.7, 0.9, 0.9, 0.9, 0.9, 1.0])
    doc.add_paragraph()

    add_heading(doc, "Period-by-period interpretation", 2)
    interpretations = [
        ("1925-1935.", " The Great Depression decade. Only 40.7 % of stocks finished positive in nominal terms "
                       "— but 48.8 % finished positive in real terms because the price level fell ~20 % over the decade."),
        ("1935-1945 through 1955-1965.", " Three consecutive decades with ~85-89 % of stocks positive. "
                                          "Post-Depression recovery, WWII industrial expansion, post-war boom."),
        ("1965-1975.", " Stagflation. 59 % positive nominal collapses to 38 % positive in real terms — "
                       "the single worst decade on record for real stock returns."),
        ("1975-1985.", " Strong recovery decade. 88 % positive nominal, 80 % positive real, median real return +133 %."),
        ("1985-1995 through 2005-2015.", " Three decades in the 55-61 % positive band. The median individual stock "
                                           "had a modest real return in the low single digits."),
        ("1985-1995 and 2015-2025.", " The most mega-cap-concentrated decades: the market returned +265 % and +252 % "
                                      "respectively, but only 15-16 % of stocks beat the market. "
                                      "Cap-weighted index returns are dominated by a small set of mega-cap names."),
    ]
    for bold_part, rest in interpretations:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    # ============= 4. Q2 30-year =============
    add_heading(doc, "4. Q2 — 30-year calendar windows", 1)

    rows30 = []
    for _, row in bench30.iterrows():
        nom = s30[(s30["window"] == row["window"]) & (s30["basis"] == "nominal")].iloc[0]
        real = s30[(s30["window"] == row["window"]) & (s30["basis"] == "real (CPI-adj)")].iloc[0]
        rows30.append([
            row["window"],
            f"{int(row['n_stocks']):,}",
            pct(nom["pct_positive"]),
            pct(real["pct_positive"]),
            signed_pct(row["vwretd_window_return"], 0),
            pct(row["pct_beat_vwretd"]),
            signed_pct(nom["median"], 1),
            signed_pct(real["median"], 1),
        ])
    add_table(doc,
              ["Window", "N", "% pos nom", "% pos real", "Market return", "% beat mkt", "Median nom", "Median real"],
              rows30,
              col_widths=[0.95, 0.7, 0.9, 0.9, 1.1, 0.9, 1.0, 1.0])
    doc.add_paragraph()

    add_picture(doc, "variation_30y.png")
    add_picture(doc, "pct_positive_30y.png")

    add_heading(doc, "Interpretation", 2)
    for bold_part, rest in [
        ("1955-1985", " is the strongest 30-year window: post-war prosperity swallowed the 1965-1975 stagflation shock "
                      "and still left ~88 % of the 1955 cohort positive."),
        ("1985-2015", " has the lowest share of stocks positive (56 %) despite the highest market total return "
                      "(+1,620 %). Only 1 in 12 stocks beat the market."),
        ("1925-1955", " has a large partial-return tail (small cohort of 508 stocks, many wiped out in the Depression)."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    # ============= 5. Q3 Full-life =============
    add_heading(doc, "5. Q3 — Full-life per stock", 1)
    add_para(doc,
             "One observation per PERMNO, compounding MthRet from the stock's first to last listed month. "
             "This is the only horizon that remains stock-anchored.")

    add_table(doc,
              ["Basis", "N", "% positive", "Median", "Mean", "Max"],
              [
                  ["Nominal", f"{int(life_nom['n']):,}", pct(life_nom["pct_positive"]),
                   signed_pct(life_nom["median"]), fmt_mult(life_nom["mean"]), fmt_mult(life_nom["max"])],
                  ["Real (CPI-adj)", f"{int(life_real['n']):,}", pct(life_real["pct_positive"]),
                   signed_pct(life_real["median"]), fmt_mult(life_real["mean"]), fmt_mult(life_real["max"])],
              ],
              col_widths=[1.5, 0.9, 1.1, 1.1, 1.1, 1.1])
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Percentiles:").bold = True
    add_table(doc,
              ["Basis", "p5", "p25", "Median", "p75", "p95"],
              [
                  ["Nominal", signed_pct(life_nom["p5"]), signed_pct(life_nom["p25"]),
                   signed_pct(life_nom["p50"]), signed_pct(life_nom["p75"]), signed_pct(life_nom["p95"])],
                  ["Real", signed_pct(life_real["p5"]), signed_pct(life_real["p25"]),
                   signed_pct(life_real["p50"]), signed_pct(life_real["p75"]), signed_pct(life_real["p95"])],
              ],
              col_widths=[1.2, 1.1, 1.1, 1.1, 1.1, 1.1])
    doc.add_paragraph()

    add_picture(doc, "hist_fulllife_nominal.png")
    add_picture(doc, "hist_fulllife_real.png")

    add_para(doc,
             "Across the entire lifespan of every US-listed common stock, the median experience is a nominal loss of ~7 % "
             "and a real loss of ~29 %. The mean is inflated by a small number of 100-year compounders — one stock returned "
             "44,211× nominal — but half the universe never crosses break-even.")

    # ============= 6. Do stocks beat the market? =============
    add_heading(doc, "6. Do stocks beat the market?", 1)
    add_para(doc, "Compounded over exactly the same months as the stock / window.")

    rows_beat = []
    for _, row in bench.iterrows():
        label = (f"{row['horizon']} {row['window']}" if row['horizon'] != 'full-life'
                 else "Full-life")
        rows_beat.append([
            label,
            f"{int(row['n_stocks']):,}",
            pct(row["pct_beat_vwretd"]),
            pct(row["pct_beat_sprtrn"]),
        ])
    add_table(doc,
              ["Horizon & window", "N", "% beat CRSP VW market", "% beat S&P 500 (price)"],
              rows_beat,
              col_widths=[2.0, 0.9, 2.0, 2.0])
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Full-span benchmarks (1925-12 to 2025-12):").bold = True
    add_table(doc,
              ["Index", "Cumulative total return", "CAGR"],
              [
                  ["CRSP value-weighted US equity (incl. dividends)",
                   f"{mkt['vwretd_total_return']*100:,.0f} % ({1+mkt['vwretd_total_return']:,.0f}×)",
                   pct(mkt['vwretd_cagr'], 2)],
                  ["S&P 500 price index (ex-dividends)",
                   f"{mkt['sprtrn_total_return']*100:,.0f} % ({1+mkt['sprtrn_total_return']:,.0f}×)",
                   pct(mkt['sprtrn_cagr'], 2)],
                  ["CPI-U (inflation)", "1,710 % (18.1×)", "2.94 %"],
              ],
              col_widths=[3.4, 2.4, 1.1])
    doc.add_paragraph()

    add_para(doc,
             "The sharp declines in '% beating the market' in 1985-1995, 2015-2025, and over the full 1985-2015 30-year "
             "window reflect the rise of extreme mega-cap dominance in the index return. The index compounds at a "
             "double-digit rate, but the share of individual stocks that keep up with it has been falling.")

    # ============= 7. Caveats =============
    add_heading(doc, "7. Caveats", 1)
    caveats = [
        "Delisting returns are not always explicitly provided. Only 3,100 of 31,565 stocks have an explicit CRSP "
        "delisting return (MthRetFlg == 'DE'); the remainder simply stop appearing at delisting. Methodology "
        "terminates the return at the last observed month. Applying the Shumway (1997) −30 % correction to "
        "performance-related delistings without explicit returns moves headline percentages by less than 1 point.",
        "Not-traded months (1.6 % of rows, flag 'NT') have MthRet set to 0 to preserve the calendar timeline.",
        "Cohort size grows over time. The 1925-1935 cohort has only 508 stocks (CRSP started with NYSE-only coverage); "
        "by 1995-2005 the cohort is 7,835 stocks (NYSE + AMEX + NASDAQ). Later-decade statistics are tighter.",
        "Windows extending beyond 2025-12 are not analyzed. The 30-year track stops at the 1985-2015 window; a 2015-2025 "
        "entry in the 30-year table would be a partial 10-year window and is excluded.",
    ]
    for i, text in enumerate(caveats, 1):
        doc.add_paragraph(f"{i}. {text}", style="List Number")

    # ============= 8. Files =============
    add_heading(doc, "8. Files produced", 1)
    files = doc.add_paragraph()
    files.add_run(
        "results/\n"
        "  report.md, report.docx\n"
        "  summary_10y.csv, summary_30y.csv, summary_fulllife.csv\n"
        "  returns_10y.csv, returns_30y.csv              — long-format per-stock returns per window\n"
        "  permno_fulllife.csv                           — per-stock full-life return\n"
        "  market_benchmark.csv, market_summary.csv\n"
        "  variation_10y.png, variation_30y.png          — box+violin across windows\n"
        "  pct_positive_10y.png, pct_positive_30y.png    — bar chart + market line\n"
        "  hist_10y_{window}_{basis}.png                 — per-window histograms (20 files)\n"
        "  hist_30y_{window}_{basis}.png                 — per-window histograms (6 files)\n"
        "  hist_fulllife_{basis}.png                     — Q3 histograms"
    ).font.name = "Consolas"
    for run in files.runs:
        run.font.size = Pt(8.5)

    # ============= 9. Reproducibility =============
    add_heading(doc, "9. Reproducibility", 1)
    p = doc.add_paragraph()
    r = p.add_run("python program/run_analysis.py    # ~50 s\npython program/build_docx.py      # ~3 s")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    add_para(doc, "Python 3.11, pandas 2.3, numpy 2.2, matplotlib 3.10, python-docx 1.2.")

    # ============= 10. References =============
    add_heading(doc, "10. References", 1)
    for ref in [
        "Bessembinder, H. (2018). \"Do stocks outperform Treasury bills?\" Journal of Financial Economics, 129(3), 440-457.",
        "Shumway, T. (1997). \"The delisting bias in CRSP data.\" Journal of Finance, 52(1), 327-340.",
        "CRSP. Monthly Stock File, schema documentation in data/Monthly Stock File.pdf.",
        "U.S. Bureau of Labor Statistics. Consumer Price Index for All Urban Consumers, CPI-U, not seasonally adjusted, monthly 1925-2025.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    build(doc)
    outpath = os.path.join(RESULTS, "report.docx")
    doc.save(outpath)
    print(f"[saved] {outpath}  ({os.path.getsize(outpath):,} bytes)")


if __name__ == "__main__":
    main()
